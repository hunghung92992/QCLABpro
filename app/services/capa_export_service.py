# -*- coding: utf-8 -*-
"""
app/services/capa_export_service.py
SERVICE XUẤT BÁO CÁO CAPA - FIXED FONT PDF (Regular & Bold)
"""
import os
import sys
from datetime import datetime

# Import thư viện an toàn
try:
    from openpyxl import load_workbook, Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from fpdf import FPDF
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError as e:
    print(f"❌ Thiếu thư viện: {e}")


class CapaExportService:
    def __init__(self):
        self.template_dir = "app/resources/templates/"
        if not os.path.exists(self.template_dir):
            os.makedirs(self.template_dir, exist_ok=True)

    def _get_val(self, data, key,
                 default="................................................................................................................................................................................"):
        if not data: return default
        val = data.get(key)
        return str(val) if val and str(val).strip() else (default + "\n" + default)

    # --- WORD HELPER ---
    def _create_element(self, name):
        return OxmlElement(name)

    def _create_attribute(self, element, name, value):
        element.set(qn(name), value)

    def _add_page_number(self, paragraph):
        run = paragraph.add_run()
        fldChar1 = self._create_element('w:fldChar')
        self._create_attribute(fldChar1, 'w:fldCharType', 'begin')
        instrText = self._create_element('w:instrText')
        self._create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"
        fldChar2 = self._create_element('w:fldChar')
        self._create_attribute(fldChar2, 'w:fldCharType', 'end')
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)

    # =========================================================================
    # 1. XUẤT WORD (.docx)
    # =========================================================================
    def export_word(self, data: dict, save_path: str):
        try:
            if os.path.exists(save_path):
                try:
                    os.rename(save_path, save_path)
                except OSError:
                    return False, "File Word đang mở. Vui lòng đóng lại!"

            doc = Document()

            # Margin
            section = doc.sections[0]
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2)

            # Header
            header = section.header
            htable = header.add_table(1, 2, width=Inches(6))
            htable.autofit = False
            htable.columns[0].width = Inches(3.5)
            htable.columns[1].width = Inches(2.5)

            h_cell1 = htable.cell(0, 0)
            p1 = h_cell1.paragraphs[0]
            p1.add_run("BỆNH VIỆN/TRUNG TÂM......\n").bold = True
            p1.add_run("KHOA XÉT NGHIỆM").bold = True

            h_cell2 = htable.cell(0, 1)
            p2 = h_cell2.paragraphs[0]
            p2.add_run("Mã hiệu: XN.QTQL.025-BM.01\n")
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Title
            title = doc.add_heading('PHIẾU KHẮC PHỤC VÀ PHÒNG NGỪA', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.color.rgb = None
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.bold = True

            capa_id = data.get('capa_id', 'SKPH-.........')
            p_ref = doc.add_paragraph(f"Số: {capa_id}")
            p_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Phần A
            doc.add_paragraph().add_run(
                "PHẦN A: Mô tả chi tiết sự không phù hợp và các biện pháp xử lý tức thời").bold = True

            doc.add_paragraph("1. Mô tả sự không phù hợp").runs[0].bold = True
            desc_text = f"Tiêu đề: {data.get('title', '')}\nChi tiết: {self._get_val(data, 'description')}"
            doc.add_paragraph(desc_text)

            doc.add_paragraph("2. Biện pháp xử lý tức thời (nếu có)").runs[0].bold = True
            doc.add_paragraph(self._get_val(data, 'correction'))

            # Footer A
            table_a = doc.add_table(rows=1, cols=2)
            table_a.autofit = True
            table_a.rows[0].cells[0].text = f"Người phát hiện: {data.get('owner', '.....................')}"

            ts = data.get('due_date', '')
            if not ts: ts = datetime.now().strftime("%d/%m/%Y")
            p_date = table_a.rows[0].cells[1].paragraphs[0]
            p_date.text = f"Ngày: {ts}"
            p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            doc.add_paragraph("\n")

            # Phần B
            doc.add_paragraph().add_run("PHẦN B: Phần dành cho cán bộ Lãnh đạo Khoa/QLCL/QLKT").bold = True

            doc.add_paragraph("1. Phân tích nguyên nhân gốc rễ").runs[0].bold = True
            doc.add_paragraph(self._get_val(data, 'root_cause'))

            p_plan = doc.add_paragraph()
            p_plan.add_run("2. Kế hoạch\t\tKhắc phục\t\t\t\tPhòng ngừa").bold = True
            doc.add_paragraph(self._get_val(data, 'corrective'))

            doc.add_paragraph("3. Thời gian thực hiện/ Người thực hiện").runs[0].bold = True
            info_exec = f"Thời gian: {self._get_val(data, 'due_date')}\nNgười thực hiện: {self._get_val(data, 'owner')}"
            doc.add_paragraph(info_exec)

            doc.add_paragraph("4. Kết luận").runs[0].bold = True
            verify_info = f"Trạng thái: {data.get('status', '')}\nMinh chứng: {self._get_val(data, 'verify_evidence')}"
            doc.add_paragraph(verify_info)

            # Ký tên
            doc.add_paragraph("\n\n")
            sig_table = doc.add_table(rows=1, cols=2)
            sig_table.rows[0].cells[0].text = "NGƯỜI LẬP PHIẾU\n(Ký và ghi rõ họ tên)"
            sig_table.rows[0].cells[1].text = "LÃNH ĐẠO KHOA / QLCL\n(Ký và ghi rõ họ tên)"
            sig_table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            sig_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Footer
            footer = section.footer
            fp = footer.paragraphs[0]
            fp.text = "BM.01/XN.QTQL.025\t\t\t\tTrang "
            self._add_page_number(fp)

            doc.save(save_path)
            return True, "Xuất file Word thành công!"
        except Exception as e:
            return False, f"Lỗi xuất Word: {str(e)}"

    # =========================================================================
    # 2. XUẤT EXCEL (.xlsx)
    # =========================================================================
    def export_excel(self, data: dict, template_name: str, output_path: str):
        try:
            if os.path.exists(output_path):
                try:
                    os.rename(output_path, output_path)
                except OSError:
                    return False, "File Excel đang mở. Vui lòng đóng lại!"

            wb = Workbook()
            ws = wb.active
            ws.title = "CAPA Report"

            header_font = Font(bold=True, color="FFFFFF", size=12)
            fill_header = PatternFill(start_color="336699", end_color="336699", fill_type="solid")
            center = Alignment(horizontal="center", vertical="center")
            thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'),
                                 bottom=Side(style='thin'))

            ws.merge_cells('A1:B1')
            ws['A1'] = f"PHIẾU KHẮC PHỤC & PHÒNG NGỪA: {data.get('capa_id', '')}"
            ws['A1'].font = Font(bold=True, size=14, color="FFFFFF")
            ws['A1'].fill = fill_header
            ws['A1'].alignment = center

            rows_data = [
                ("PHẦN A: MÔ TẢ & XỬ LÝ TỨC THỜI", ""),
                ("1. Mô tả sự không phù hợp", f"{data.get('title', '')}\n{data.get('description', '')}"),
                ("2. Xử lý tức thời", data.get("correction")),
                ("Người phát hiện", data.get("owner")),
                ("PHẦN B: PHÂN TÍCH & HÀNH ĐỘNG", ""),
                ("1. Nguyên nhân gốc rễ", data.get("root_cause")),
                ("2. Kế hoạch", data.get("corrective")),
                ("3. Thời gian / Người thực hiện", f"{data.get('due_date')} / {data.get('owner')}"),
                ("4. Kết luận", data.get("verify_evidence")),
                ("Trạng thái", data.get("status"))
            ]

            curr_row = 2
            for label, val in rows_data:
                cell_lbl = ws.cell(row=curr_row, column=1, value=label)
                cell_val = ws.cell(row=curr_row, column=2, value=str(val) if val else "")
                cell_lbl.border = thin_border
                cell_val.border = thin_border

                if not val and "PHẦN" in label:
                    ws.merge_cells(start_row=curr_row, start_column=1, end_row=curr_row, end_column=2)
                    cell_lbl.alignment = center
                    cell_lbl.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")

                curr_row += 1

            ws.column_dimensions['A'].width = 30
            ws.column_dimensions['B'].width = 70
            wb.save(output_path)
            return True, "Xuất Excel thành công!"
        except Exception as e:
            return False, f"Lỗi Excel: {str(e)}"

    # =========================================================================
    # 3. XUẤT PDF (.pdf) - SỬA LỖI FONT
    # =========================================================================
    def export_pdf(self, data: dict, output_path: str):
        try:
            if os.path.exists(output_path):
                try:
                    os.rename(output_path, output_path)
                except OSError:
                    return False, "File PDF đang mở. Vui lòng đóng lại!"

            pdf = FPDF()

            # --- XỬ LÝ FONT (FIX LỖI Undefined Font) ---
            windows_dir = os.environ.get('WINDIR', 'C:' + '\\Windows')
            path_reg = os.path.join(windows_dir, 'Fonts', 'arial.ttf')
            path_bold = os.path.join(windows_dir, 'Fonts', 'arialbd.ttf')

            font_name = "Arial"
            has_bold = False

            # Nạp Font Regular
            if os.path.exists(path_reg):
                try:
                    pdf.add_font('ArialUni', '', path_reg, uni=True)
                    font_name = 'ArialUni'
                except:
                    pass

            # Nạp Font Bold (Quan trọng để dùng style='B')
            if os.path.exists(path_bold) and font_name == 'ArialUni':
                try:
                    pdf.add_font('ArialUni', 'B', path_bold, uni=True)
                    has_bold = True
                except:
                    pass

            # Hàm set font an toàn (Tránh crash nếu không có bold)
            def safe_set_font(style='', size=11):
                f_style = style
                if 'B' in style and not has_bold and font_name == 'ArialUni':
                    f_style = style.replace('B', '')  # Bỏ in đậm nếu không có file font đậm
                pdf.set_font(font_name, f_style, size)

            pdf.add_page()
            safe_set_font('', 10)

            # Header giả lập
            pdf.cell(90, 10, "KHOA XÉT NGHIỆM", 1, 0, 'C')
            pdf.cell(0, 10, "Mã hiệu: XN.QTQL.025-BM.01", 1, 1, 'C')
            pdf.ln(5)

            # Title
            safe_set_font('B', 14)
            pdf.cell(0, 10, "PHIẾU KHẮC PHỤC VÀ PHÒNG NGỪA", ln=True, align='C')

            safe_set_font('', 11)
            pdf.cell(0, 8, f"Số: {data.get('capa_id', '...........')}", ln=True, align='C')
            pdf.ln(5)

            # Hàm in (An toàn font)
            def safe_print(title, content):
                safe_set_font('B', 11)
                pdf.cell(0, 6, title, ln=True)
                safe_set_font('', 11)

                txt = str(content) if content and str(
                    content).strip() else ".................................................................................................................\n................................................................................................................."
                pdf.multi_cell(0, 6, txt)
                pdf.ln(2)

            # Phần A
            safe_set_font('B', 12)
            pdf.cell(0, 10, "PHẦN A: Mô tả chi tiết & Biện pháp xử lý tức thời", ln=True)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(2)

            safe_print("1. Mô tả sự không phù hợp:", f"{data.get('title', '')}\n{data.get('description', '')}")
            safe_print("2. Biện pháp xử lý tức thời:", data.get('correction'))

            pdf.ln(2)
            pdf.cell(100, 6, f"Người phát hiện: {data.get('owner', '................')}")
            pdf.cell(0, 6, f"Ngày: {data.get('due_date', '..../..../....')}", ln=True, align='R')
            pdf.ln(5)

            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(5)

            # Phần B
            safe_set_font('B', 12)
            pdf.cell(0, 10, "PHẦN B: Dành cho Lãnh đạo Khoa/QLCL", ln=True)

            safe_print("1. Phân tích nguyên nhân gốc rễ:", data.get('root_cause'))
            safe_print("2. Kế hoạch (Khắc phục & Phòng ngừa):", data.get('corrective'))
            safe_print("3. Thời gian & Người thực hiện:", f"{data.get('due_date', '')} - {data.get('owner', '')}")
            safe_print("4. Kết luận:", data.get('verify_evidence'))

            # Footer
            pdf.ln(10)
            if pdf.get_y() > 240: pdf.add_page()

            pdf.cell(90, 6, "NGƯỜI LẬP PHIẾU", align='C')
            pdf.cell(90, 6, "PHÊ DUYỆT (LĐ KHOA)", align='C', ln=True)
            pdf.ln(15)
            safe_set_font('', 9)
            pdf.cell(90, 6, "(Ký và ghi rõ họ tên)", align='C')
            pdf.cell(90, 6, "(Ký và ghi rõ họ tên)", align='C')

            pdf.output(output_path)
            return True, "Xuất PDF thành công!"

        except Exception as e:
            return False, f"Lỗi xuất PDF: {str(e)}"